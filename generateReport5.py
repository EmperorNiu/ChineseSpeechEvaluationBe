from docx import Document
from docx import table
import sys
from pypinyin import pinyin
from sqlalchemy import Column,String,Integer,ForeignKey,create_engine,PrimaryKeyConstraint
from sqlalchemy.orm import sessionmaker
from sqlalchemy.ext.declarative import declarative_base
from collections import Counter
# doc_path = sys.argv[1]
student_id = sys.argv[1]
homework_doc_id = sys.argv[2]
student_name = sys.argv[3]
teacher = sys.argv[4]
title = sys.argv[5]
time = sys.argv[6]
audio = sys.argv[7]
doc_path = './report_template3.docx'

# student_id = 'PS001'
# homework_doc_id = 40

Base = declarative_base()
class HomeworkDoc(Base):
    __tablename__ = 'homework_doc'
    homework_doc_id = Column(Integer,primary_key=True) 
    title = Column(String)
    describe = Column(String)
    position = Column(String)
class HomeworkResult(Base):
    __tablename__ = 'student_homework_result'
    student_homework_result_id = Column(Integer,primary_key=True)
    homework_doc_id_refer = Column(Integer)
    student_id_refer = Column(String)
    tone_accuracy = Column(Integer)
    intonation_accuracy = Column(Integer)
    fluency = Column(Integer)
    comment = Column(String)
    is_thesis_express = Column(Integer)
    minus_word_error_score = Column(Integer)
class WordErr(Base):
    __tablename__ = 'word_error'
    word_error_id = Column(Integer,primary_key=True)
    student_homework_result_id_refer = Column(Integer)
    studend_id_refer = Column(String)
    word = Column(String)
    error_types = Column(String)
    whole_word = Column(String)
class WordDict(Base):
    __tablename__ = 'bigcharstock'
    index = Column(Integer,primary_key=True)
    word = Column(String)
    type = Column(String)
    pinyin = Column(String)
    shengmu = Column(String)
    yunmu1 = Column(String)
    zhuyuanyin = Column(String)
    tone = Column(Integer)
    yunmu2 = Column(String)

engine = create_engine('mysql+mysqlconnector://root:123456@localhost:3306/charword')
DBSession = sessionmaker(bind=engine)
db = DBSession()

re = db.query(HomeworkResult).filter_by(homework_doc_id_refer=homework_doc_id,student_id_refer=student_id).all()[0]
hdoc = db.query(HomeworkDoc).filter_by(homework_doc_id=homework_doc_id).all()[0]
errs = db.query(WordErr).filter_by(student_homework_result_id_refer=re.student_homework_result_id).all()

err_result = []
for err in  errs:
    tmp = err.error_types.split(',')
    err_result.extend(tmp)
err_result = [i for i in err_result if i != '']
c = Counter(err_result)
# print(c['A'])

# doc = Document(hdoc.position)
doc = Document(doc_path)

doc.tables[0].rows[7].cells[11].text = str(int(re.tone_accuracy*0.2))
doc.tables[0].rows[7].cells[12].text = str(int(re.intonation_accuracy*0.1))
doc.tables[0].rows[7].cells[13].text = str(int(re.fluency*0.1))

doc.tables[0].rows[7].cells[0].text = str(c['A'])
doc.tables[0].rows[7].cells[1].text = str(c['B'])
doc.tables[0].rows[7].cells[2].text = str(c['C'])
doc.tables[0].rows[7].cells[3].text = str(c['D'])
doc.tables[0].rows[7].cells[5].text = str(c['E'])
doc.tables[0].rows[7].cells[7].text = str(c['F'])
doc.tables[0].rows[7].cells[8].text = str(c['G'])
doc.tables[0].rows[7].cells[9].text = str(c['H'])
doc.tables[0].rows[7].cells[10].text = str(c['I'])

doc.tables[0].rows[22].cells[0].text = '评语：\n' + re.comment

tmp = 60 - c['A'] - c['B'] - c['C'] - c['D'] - c['E'] - c['F'] - c['G'] - c['H'] - c['I']
if tmp < 0:
    tmp = 0
doc.tables[0].rows[0].cells[6].text = title
doc.tables[0].rows[2].cells[6].text = student_name
doc.tables[0].rows[3].cells[6].text = teacher
doc.tables[0].rows[2].cells[12].text = student_id
doc.tables[0].rows[1].cells[6].text = time
doc.tables[0].rows[1].cells[12].text = time
doc.tables[0].rows[0].cells[12].text = audio
s1 = int(re.tone_accuracy*0.2)
s2 = int(re.intonation_accuracy*0.1)
s3 = int(re.fluency*0.1)
score = tmp + s1 + s2 + s3
doc.tables[0].rows[3].cells[12].text = str(int(score))
err_record = {}
for err in errs:
    if err.error_types != '':
        if err.word in err_record:
            err_record[err.word] += ','
            err_record[err.word] += err.error_types
        else:
            err_record[err.word] = err.error_types
# print(err_record)

err_to_word = {
    'E1': [],
    'E2': [],
    'E3': [],
    'E4': [],
    'F1': [],
    'F2': [],
    'C': [],
    'D': [],
    'G': [],
    'H': [],
    'A': [],
    'B': [],
    'I': []
}
for err in errs:
    errType = err.error_types.split(',')
    # print(err.word)
    try:
        wordDic = db.query(WordDict).filter_by(word=err.word).all()[0]
    except IndexError:
        continue
    else: 
        if 'A' in errType:
            err_to_word['A'].append(err.word)
        if 'B' in errType:
            err_to_word['B'].append(err.word)
        if 'C' in errType:
            err_to_word['C'].append(err.word + '(' + wordDic.shengmu + ')')
        if 'D' in errType:
            err_to_word['D'].append(err.word + '(' + wordDic.yunmu2 + ')')
        if 'E' in errType:
            if wordDic.tone == 1:
                err_to_word['E1'].append(err.word)
            if wordDic.tone == 2:
                err_to_word['E2'].append(err.word)        
            if wordDic.tone == 3:
                err_to_word['E3'].append(err.word)
            if wordDic.tone == 4:
                err_to_word['E4'].append(err.word)
        if 'G' in errType:
            err_to_word['G'].append(err.whole_word)
        if 'H' in errType:
            err_to_word['H'].append(err.whole_word)
        if 'F' in errType:
            dics = []
            flag = 0
            for ch in err.whole_word:
                try:
                    tmp = db.query(WordDict).filter_by(word=ch).all()[0]
                except IndexError:
                    flag = 0
                if tmp.tone == 3:
                    flag = 1
                else:
                    flag = 0
            if flag == 1:
                err_to_word['F2'].append(err.whole_word)
            else:
                err_to_word['F1'].append(err.whole_word)
        if 'I' in errType:
            tmp = pinyin(err.whole_word)
            tmp2 = []
            for i in tmp:
                tmp2.append(i[0])
            tmp = ' '.join(tmp2)
            err_to_word['I'].append(err.whole_word + '(' + tmp + ')')

# print(err_to_word)     

# 写doc：错误类型统计
doc.tables[0].rows[9].cells[4].text = str(len(err_to_word['E1']))
doc.tables[0].rows[10].cells[4].text = str(len(err_to_word['E2']))
doc.tables[0].rows[11].cells[4].text = str(len(err_to_word['E3']))
doc.tables[0].rows[12].cells[4].text = str(len(err_to_word['E4']))
doc.tables[0].rows[13].cells[4].text = str(len(err_to_word['F1']))
doc.tables[0].rows[14].cells[4].text = str(len(err_to_word['F2']))
doc.tables[0].rows[15].cells[4].text = str(len(err_to_word['C']))
doc.tables[0].rows[16].cells[4].text = str(len(err_to_word['D']))
doc.tables[0].rows[17].cells[4].text = str(len(err_to_word['G']))
doc.tables[0].rows[18].cells[4].text = str(len(err_to_word['H']))
doc.tables[0].rows[19].cells[4].text = str(len(err_to_word['A']))
doc.tables[0].rows[20].cells[4].text = str(len(err_to_word['B']))
doc.tables[0].rows[21].cells[4].text = str(len(err_to_word['I']))
# 写doc：错误细节
doc.tables[0].rows[9].cells[11].text = ','.join(err_to_word['E1'])
doc.tables[0].rows[10].cells[11].text = ','.join(err_to_word['E2'])
doc.tables[0].rows[11].cells[11].text = ','.join(err_to_word['E3'])
doc.tables[0].rows[12].cells[11].text = ','.join(err_to_word['E4'])
doc.tables[0].rows[13].cells[11].text = ','.join(err_to_word['F1'])
doc.tables[0].rows[14].cells[11].text = ','.join(err_to_word['F2'])
doc.tables[0].rows[15].cells[11].text = ','.join(err_to_word['C'])
doc.tables[0].rows[16].cells[11].text = ','.join(err_to_word['D'])
doc.tables[0].rows[17].cells[11].text = ','.join(err_to_word['G'])
doc.tables[0].rows[18].cells[11].text = ','.join(err_to_word['H'])
doc.tables[0].rows[19].cells[11].text = ','.join(err_to_word['A'])
doc.tables[0].rows[20].cells[11].text = ','.join(err_to_word['B'])
doc.tables[0].rows[21].cells[11].text = ','.join(err_to_word['I'])
# err_str = '错误记录：\n'
# for key in err_record:
#     err_str += (key + ':\t' + err_record[key] + '\n')
# doc.tables[0].rows[8].cells[0].text = err_str
# doc.tables[0].rows[9].cells[0].text = ''

dname = './download/doc/report_' + student_id + '_' + str(homework_doc_id) + '.docx'
doc.save(dname)
print("success")
db.close()
# tmp_name = './report_' + student_id + '_' + str(homework_doc_id) + '.docx'
# doc.save(tmp_name)
