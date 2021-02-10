from docx import Document
from docx import table
import sys
from sqlalchemy import Column,String,Integer,ForeignKey,create_engine,PrimaryKeyConstraint
from sqlalchemy.orm import sessionmaker
from sqlalchemy.ext.declarative import declarative_base
from collections import Counter
doc_path = sys.argv[1]
student_id = sys.argv[1]
homework_doc_id = sys.argv[2]
student_name = sys.argv[3]
teacher = sys.argv[4]
title = sys.argv[5]
time = sys.argv[6]
audio = sys.argv[7]
# doc_path = '../data/诊断报告.docx'
# student_id = 'PS005'
# homework_doc_id = 1
Base = declarative_base()
class HomeworkDoc(Base):
    __tablename__ = 'homework_doc'
    homework_doc_id = Column(Integer,primary_key=True) 
    title = Column(String)
    describe = Column(String)
    position = Column(String)
class HomeworkResult(Base):
    __tablename__ = 'student_homework_result'
    student_homework_result_id = Column(Integer,primary_key=True)
    homework_doc_id_refer = Column(Integer)
    student_id_refer = Column(String)
    tone_accuracy = Column(Integer)
    intonation_accuracy = Column(Integer)
    fluency = Column(Integer)
    comment = Column(String)
class WordErr(Base):
    __tablename__ = 'word_error'
    word_error_id = Column(Integer,primary_key=True)
    student_homework_result_id_refer = Column(Integer)
    studend_id_refer = Column(String)
    word = Column(String)
    error_types = Column(String)

engine = create_engine('mysql+mysqlconnector://root:123456@localhost:3306/charword')
DBSession = sessionmaker(bind=engine)
db = DBSession()

re = db.query(HomeworkResult).filter_by(homework_doc_id_refer=homework_doc_id,student_id_refer=student_id).all()[0]
hdoc = db.query(HomeworkDoc).filter_by(homework_doc_id=homework_doc_id).all()[0]
errs = db.query(WordErr).filter_by(student_homework_result_id_refer=re.student_homework_result_id).all()

err_result = []
for err in  errs:
    tmp = err.error_types.split(',')
    err_result.extend(tmp)
err_result = [i for i in err_result if i != '']
c = Counter(err_result)
# print(c['A'])

doc = Document(hdoc.position)
# doc = Document(doc_path)
doc.tables[0].rows[7].cells[9].text = str(int(re.tone_accuracy*0.2))
doc.tables[0].rows[7].cells[10].text = str(int(re.intonation_accuracy*0.1))
doc.tables[0].rows[7].cells[12].text = str(int(re.intonation_accuracy*0.1))

doc.tables[0].rows[7].cells[0].text = str(c['A'])
doc.tables[0].rows[7].cells[1].text = str(c['B'])
doc.tables[0].rows[7].cells[2].text = str(c['C'])
doc.tables[0].rows[7].cells[3].text = str(c['D'])
doc.tables[0].rows[7].cells[5].text = str(c['E'])
doc.tables[0].rows[7].cells[6].text = str(c['F'])
doc.tables[0].rows[7].cells[7].text = str(c['G'])
doc.tables[0].rows[7].cells[8].text = str(c['H'])
doc.tables[0].rows[4].cells[13].text = '评语：\n' + re.comment

tmp = 60 - c['A'] - c['B'] - c['C'] - c['D'] - c['E'] - c['F'] - c['G'] - c['H']
if tmp < 0:
    tmp = 0
doc.tables[0].rows[0].cells[5].text = title
doc.tables[0].rows[2].cells[5].text = student_name
doc.tables[0].rows[3].cells[5].text = teacher
doc.tables[0].rows[2].cells[11].text = student_id
doc.tables[0].rows[1].cells[5].text = time
doc.tables[0].rows[1].cells[11].text = time
doc.tables[0].rows[0].cells[11].text = audio
score = tmp + re.tone_accuracy*0.2 +re.intonation_accuracy*0.1 + re.intonation_accuracy*0.1
doc.tables[0].rows[3].cells[11].text = str(score)
dname = './download/doc/report_' + student_id + '_' + str(homework_doc_id) + '.docx'
doc.save(dname)

print(dname)