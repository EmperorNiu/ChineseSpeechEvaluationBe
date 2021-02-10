from docx import Document
from docx import table
import re
import pandas as pd
import sys
from sqlalchemy import Column,String,Integer,ForeignKey,create_engine,PrimaryKeyConstraint
from sqlalchemy.orm import sessionmaker
from sqlalchemy.ext.declarative import declarative_base

Base = declarative_base()
class Article(Base):
    __tablename__ = 'homework_article'
    article_id = Column(Integer,primary_key=True,autoincrement = True)
    content = Column(String(1000))
    homework_index = Column(Integer)
class Homework(Base):
    __tablename__ = 'homework'
    exercise_id = Column(Integer,primary_key=True,autoincrement = True)
    character = Column(String(5))
    word1 = Column(String(10))
    word2 = Column(String(10))
    homework_index = Column(Integer)

doc_path = sys.argv[1]
homework_doc_id = sys.argv[2]

doc = Document(doc_path)
# doc = Document('../data/诊断报告表头样式.docx')

print(doc.paragraphs[0].text)
content1 = doc.tables[0].rows[8].cells[0].text
content2 = doc.tables[0].rows[9].cells[0].text

a = content1.split('\t')
pattern = r'[\n\t]'
a = re.split(pattern,content1)
a = [i for i in a if i != '']
b = [i.split(' ') for i in a][1:]


engine = create_engine('mysql+mysqlconnector://root:123456@localhost:3306/charword')

DBSession = sessionmaker(bind=engine)
db = DBSession()
Base.metadata.create_all(engine)
article = Article(content = content2,homework_index=homework_doc_id)
db.add(article)
db.commit()
for i in b:
    if(len(i)==0):
        continue
    hw = Homework(character=i[0],word1=i[1],homework_index=homework_doc_id)
    if len(i) == 3:
        hw.word2 = i[2]
    db.add(hw)
db.commit()
db.close()